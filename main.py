from Travelling.bus import Bus
from Travelling.train import Train
from Travelling.plane import Plane
from docx import Document
from openpyxl import Workbook
import sqlite3

"""АВТОБУС"""
def traveling_by_bus():
    print("Путешесвтие на автобусе: \n")
    try:
        distance = float(input("Введите расстоение маршрута в км: "))
        passengers = int(input("Введите количество пассажиров: "))
        price = float(input("Введите стоимость билета: "))

        if (distance >= 0 and passengers >= 0 and price >= 0):
            # создание экземпляра класса
            bus = Bus(distance, passengers, price)
            print(f"\nСтоимость маршрута: {bus.calculate_cost()} руб.")
            print(f"Время поездки: {bus.calculate_time()} часов" )
            print(f"Пассажиропоток: {bus.calculate_passenger_flow()}" )
            choosing_to_save(bus, "Автобус")
        else:
            print("Ошибка! Параметры должны быть положительными числами")
    except(ValueError):
        print("Ошибка! Неверный ввод параметров")

""" ПОЕЗД """
def traveling_by_train():
    print("Путешествие на поезде: \n")
    try:
        distance = float(input("Введите расстоение маршрута в км: "))
        passengers = int(input("Введите количество пассажиров: "))
        price = float(input("Введите стоимость билета: "))

        if (distance >= 0 and passengers >= 0 and price >= 0):
            #создание экземпляра класса
            train = Train(distance, passengers, price)
            print(f"\nСтоимость маршрута: {train.calculate_cost()} руб.")
            print(f"Время поездки: {train.calculate_time()} часов")
            print(f"Пассажиропоток: {train.calculate_passenger_flow()}")
            choosing_to_save(train, "Поезд")
        else:
            print("Ошибка! Параметры должны быть положительными числами")
    except(ValueError):
        print("Ошибка! Неверный ввод параметров")

""" САМОЛЕТ"""
def traveling_by_plane():
    print("Путешествие на самолете: \n")
    try:
        distance = float(input("Введите расстоение маршрута в км: "))
        passengers = int(input("Введите количество пассажиров: "))
        price = float(input("Введите стоимость билета: "))

        if (distance >= 0 and passengers >= 0 and price >= 0):
            # создание экземпляра класса
            plane = Plane(distance, passengers, price)
            print(f"\nСтоимость маршрута: {plane.calculate_cost()} руб.")
            print(f"Время поездки: {plane.calculate_time()} часов")
            print(f"Пассажиропоток: {plane.calculate_passenger_flow()}")
            choosing_to_save(plane, "Самолет")
        else:
            print("Ошибка! Параметры должны быть положительными числами")
    except(ValueError):
        print("Ошибка! Неверный ввод параметров")

"""Сохранение в  word"""
def word(transport_type, file_name):
    doc = Document()
    doc.title = f"Отчет <<Путешествие на {transport_type}>>"
    doc.add_heading("Характеристики маршрута", 2)
    doc.add_paragraph(f"Расстояние маршрута: {transport_type.distance} км")
    doc.add_paragraph(f"Количество пассажиров: {transport_type.passengers}  ")
    doc.add_paragraph(f"Стоимость билета: {transport_type.ticket_price} руб.")

    doc.add_paragraph()

    doc.add_heading("Расчеты: ", 2)

    doc.add_paragraph(f"Стоимость маршрута: {transport_type.distance} руб.")
    doc.add_paragraph(f"Время поездки: {transport_type.distance} час.")
    doc.add_paragraph(f"Пассажиропоток: {transport_type.distance} ")

    doc.save(f"{file_name}.docx")
    print("Отчет сохранен")

"""Сохранение в excel"""
def excel(transport_obj, transport_type):
    wb = Workbook()
    ws = wb.active

    ws['A1'] = f"Отчет <<Путешествие на {transport_type}>>"
    ws['A3'] = "Характеристики маршрута"

    parameters = [
        ('Расстояние маршрута', f'{transport_obj.distance} км'),
        ('Количество пассажиров', transport_obj.passengers),
        ('Стоимость билета', f'{transport_obj.ticket_price} руб.')
    ]
    for i, (param, value) in enumerate(parameters, 4):
        ws[f'A{i}'] = param
        ws[f'B{i}'] = value

    ws['A7'] = "Расчеты:"
    results = [
        ('Общая стоимость маршрута', f'{transport_obj.calculate_cost()}'),
        ('Время поездки', f'{transport_obj.calculate_time()}'),
        ('Пассажиропоток', f'{transport_obj.calculate_passenger_flow()} ')
    ]
    for i, (result, value) in enumerate(results, 10):
        ws[f'A{i}'] = result
        ws[f'B{i}'] = value
    wb.save(f"{transport_type}.xlsx")
    print("Отчет сохранен")

"""сохранение в БД"""
def init_db():
    con = sqlite3.connect("traveling.db")
    cursor = con.cursor()
    cursor.execute('''
                CREATE TABLE IF NOT EXISTS transport_calculations (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    transport_type TEXT NOT NULL,
                    distance REAL NOT NULL,
                    passengers INTEGER NOT NULL,
                    ticket_price REAL NOT NULL,
                    total_cost REAL NOT NULL,
                    travel_time REAL NOT NULL,
                    passenger_flow REAL NOT NULL
                )
            ''')
    cursor.execute('''
                CREATE TABLE IF NOT EXISTS reports (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    report_type TEXT NOT NULL,
                    transport_type TEXT NOT NULL,
                    filename TEXT NOT NULL
                )
            ''')

    con.commit()
    con.close()

def sqlite(transport_obj, transport_type):
    try:
        con = sqlite3.connect("traveling.db")
        cursor = con.cursor()
        cursor.execute('''
                    INSERT INTO transport_calculations 
                    (transport_type, distance, passengers, ticket_price, total_cost, travel_time, passenger_flow)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                ''', (
            transport_type,
            transport_obj.distance,
            transport_obj.passengers,
            transport_obj.ticket_price,
            transport_obj.calculate_cost(),
            transport_obj.calculate_time(),
            transport_obj.calculate_passenger_flow()
        ))

        con.commit()
        con.close()

        print("Отчет сохранен")
    except(ValueError):
        print("Ошибка!")
def choosing_to_save(transport_obj, transport_type):
    try:
        save = input("Сохранить отчет? (yes/no)")
        if (save == "yes" or save == "Yes" or save == "да" or save == "Да"):
            print("\nВыберите формат отчета: ")
            print("1. Word")
            print("2. Excel")
            print("3. SQLite")
            print("4. Закрыть")
            while True:
                try:
                    num = int(input("Ваш выбор: "))

                    if num == 1:
                        word(transport_obj, transport_type)
                    elif num == 2:
                        excel(transport_obj, transport_type)
                    elif num == 3:
                        sqlite(transport_obj, transport_type)
                    elif num == 4:
                        break
                    else:
                        print("Выберите действие от 1 до 4")
                except(ValueError):
                    print("Ошибка! Неверный ввод ")
    except(ValueError):
        print("Ошибка! Неверный ввод")

def main():
    init_db()
    while(True):
        print("Выберите действие: \n")
        print("1. Путешествие на автобусе")
        print("2. Путешесвтие на поезде")
        print("3. Путешествие на самолете")
        print("4. Выход")

        try:
            choice = int(input("Ваш выбор: "))
            if (choice < 1 or choice > 4):
                print("Ошибка! Выберите действие (1-4)")
            if choice == 1:
                traveling_by_bus()
            elif choice == 2:
                traveling_by_train()
            elif choice == 3:
                traveling_by_plane()
            elif choice == 4:
                break
        except(ValueError):
            print("Ошибка! Неверный ввод")

if __name__ == '__main__':
    main()